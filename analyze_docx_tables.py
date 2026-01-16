
from docx import Document
import sys

path = r"d:\HACC_DEV\rapport mensuel\REPORTING ADMINISTRATIF-FINANCIER-VENTES JUIN 2025.docx"
try:
    doc = Document(path)
    print(f"Total tables: {len(doc.tables)}")

    for i, table in enumerate(doc.tables):
        print(f"\n--- Table {i} ---")
        # Print first 5 rows
        for j, row in enumerate(table.rows[:5]):
            row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(row_text)
except Exception as e:
    print(f"Error: {e}")
